from datetime import datetime
import time
import baostock as bs
import matplotlib.pyplot as plt
import docx
from docx.shared import RGBColor
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

from tools.config import MyMysql


# 获取股票今日金融新闻
def get_today_news(code):
    database = MyMysql()
    db = database.connect
    cur = database.cursor

    # 返回一个状态码，成功获取今日新闻置1，无今日但获取3日内新闻置2，未成功获取新闻置0
    status = 0

    sql = 'select content, datetime, score ' \
          'from sina_finance ' \
          'where TO_DAYS(datetime) = TO_DAYS(NOW()) and code = %s'
    cur.execute(sql, (code,))
    data = cur.fetchall()
    if data:
        status = 1
    else:
        sql = 'select content, datetime, score ' \
              'from sina_finance ' \
              'where TO_DAYS(NOW()) - TO_DAYS(datetime) <= 3 and code = %s'
        cur.execute(sql, (code,))
        data = cur.fetchall()
        if data:
            status = 2

    cur.close()
    db.close()
    return status, data


# 获取股票今日股吧舆论
def get_today_guba(code):
    database = MyMysql()
    db = database.connect
    cur = database.cursor

    # 转换股票代码为东财股吧识别格式
    if code == 'sh000001':
        code = 'zssh000001'
    elif 'sh' in code:
        code = code.replace('sh', '')

    # 返回一个状态码，成功获取今日舆论置1，未成功获取舆论置0
    status = 0

    sql = 'select content, author, view_num, comment_num, score, composite_score, datetime ' \
          'from eastmoney_guba ' \
          'where TO_DAYS(datetime) = TO_DAYS(NOW()) and code = %s ' \
          'order by comment_num desc '
    cur.execute(sql, (code,))
    data = cur.fetchall()
    if data:
        status = 1

    cur.close()
    db.close()
    return status, data


# 根据股票代码获取股票名称
def code_to_name(code):
    code = code[:2] + '.' + code[2:]
    # 登陆系统
    bs.login()
    # 获取证券基本资料
    rs = bs.query_stock_basic(code=code)
    # 打印结果集
    data_list = []
    while (rs.error_code == '0') & rs.next():
        # 获取一条记录，将记录合并在一起
        data_list.append(rs.get_row_data())
    # 登出系统
    bs.logout()

    return data_list[0][1]


# 生成doc格式报告
def generate_report(code):
    code_name = code_to_name(code)

    doc = docx.Document()
    # 设置字体(正文）
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 创建分析报告封面
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_before = Pt(200)
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run(code_name + '舆情分析报告')
    run.font.color.rgb = RGBColor(60, 60, 60)
    run.font.bold = True
    run.font.size = Pt(36)

    date = time.strftime("%Y年%m月%d日", time.localtime())
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(date)
    run.font.color.rgb = RGBColor(105, 105, 105)
    run.font.size = Pt(22)

    doc.add_page_break()
    # 创建分析报告正文
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('一、股票相关新闻数据')
    run.font.color.rgb = RGBColor(60, 60, 60)
    run.font.size = Pt(20)
    run.font.bold = True
    # 获取并显示新闻内容
    status, news = get_today_news(code)
    news_num = len(news)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
    p.paragraph_format.first_line_indent = Inches(0.32)  # 首行缩进
    if status == 0:
        introduction = '今日共有' + code_name + '相关新闻内容0篇，且近三日内无相关新闻。'
    elif status == 1:
        introduction = '今日共有' + code_name + '相关新闻内容' + str(news_num) + '篇, 数据来源为新浪财经，具体如下：'
    else:
        introduction = '今日共有' + code_name + '相关新闻内容0篇, 近三日内共有相关新闻' + str(news_num) + '篇，数据来源为新浪财经，具体如下：'
    run = p.add_run(introduction)
    run.font.size = Pt(14)
    # 以表格形式展示新闻
    tb = doc.add_table(news_num+1, 3)
    tb.style = 'Medium Grid 3 Accent 1'
    tb.style.font.size = Pt(12)
    tb.cell(0, 1).width = Inches(8)
    heading_cells = tb.rows[0].cells
    heading_cells[0].text = '监控公司'
    heading_cells[1].text = '新闻标题'
    heading_cells[2].text = '发布时间'
    for i in range(news_num):
        tb.cell(i + 1, 0).text = code_name
        tb.cell(i + 1, 1).text = news[i][0]
        tb.cell(i + 1, 2).text = datetime.strftime(news[i][1], '%Y-%m-%d %H:%M:%S')
    # 展示新闻情感分析结果
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
    p.paragraph_format.first_line_indent = Inches(0.32)  # 首行缩进
    run = p.add_run('对上述新闻作情感分析，得到正面与负面新闻占比如下图所示：')
    run.font.size = Pt(14)
    # 使用matplotlib绘制正负面新闻占比饼状图
    pos_num = 0
    neg_num = 0
    neu_num = 0
    for per in news:
        score = per[2]
        if score > 0:
            pos_num += 1
        elif score < 0:
            neg_num += 1
        else:
            neu_num += 1
    plt.rcParams['font.family'] = ['Heiti TC']  # 设置正常显示中文
    plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号
    labels = ['积极', '中性', '消极']  # 定义饼的标签名称
    fraces = [pos_num, neu_num, neg_num]  # 显示百分比数据
    max_num = pos_num
    max_att = '积极'
    if neu_num > max_num:
        max_num = neu_num
        max_att = '中性'
    if neg_num > max_num:
        max_att = '消极'
    explode = [0.1, 0, 0]  # 饼图分离
    plt.axis('equal')  # 设置x,y的刻度一样，使其饼图为正圆
    plt.figure(figsize=(6, 2.5), dpi=150)
    cmap = plt.get_cmap("tab20c")  # 选择tab20c的色彩，
    color = cmap([6, 9, 1])  # 选择tab20c中的4种颜色
    plt.pie(x=fraces, labels=labels, autopct='%0.2f%%', colors=color, explode=explode, shadow=True)  # 绘制饼状图
    plt.legend()
    plt.savefig('../docs/pics/guba_pie.jpg')  # 保存到本地文件夹，当前路径下
    doc.add_picture('../docs/pics/guba_pie.jpg')  # 向报告中插入图片
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.32)  # 首行缩进
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run('其中，积极倾向的舆论共' + str(pos_num) + '条，中性倾向的舆论共' + str(neu_num) +
                    '条， 消极倾向的舆论共' + str(neg_num) + '条。舆论总体偏向' + max_att + '。')
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('二、股票相关舆论数据')
    run.font.color.rgb = RGBColor(60, 60, 60)
    run.font.size = Pt(20)
    run.font.bold = True
    # 获取并显示舆论内容
    status, guba = get_today_guba(code)
    guba_num = len(guba)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
    p.paragraph_format.first_line_indent = Inches(0.32)  # 首行缩进
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(14)
    if status == 0:
        introduction = '今日共有' + code_name + '相关舆情内容0篇，数据来源为东方财富股吧。'
    else:
        introduction = '今日共有' + code_name + '相关舆情内容' + str(guba_num) + '篇, 数据来源为东方财富股吧，分析如下：'
    run = p.add_run(introduction)
    run.font.size = Pt(14)
    # 使用matplotlib绘制正负面新闻占比饼状图
    pos_num = 0
    neg_num = 0
    neu_num = 0
    for per in guba:
        score = per[4]
        if score > 0:
            pos_num += 1
        elif score < 0:
            neg_num += 1
        else:
            neu_num += 1
    plt.rcParams['font.family'] = ['Heiti TC']  # 设置正常显示中文
    plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号
    labels = ['积极', '中性', '消极']  # 定义饼的标签名称
    fraces = [pos_num, neu_num, neg_num]  # 显示百分比数据
    explode = [0.1, 0, 0]  # 饼图分离
    plt.axis('equal')  # 设置x,y的刻度一样，使其饼图为正圆
    plt.figure(figsize=(6, 2.5), dpi=150)
    cmap = plt.get_cmap("tab20c")  # 选择tab20c的色彩，
    color = cmap([6, 9, 1])  # 选择tab20c中的4种颜色
    plt.pie(x=fraces, labels=labels, autopct='%0.2f%%', colors=color, explode=explode, shadow=True)  # 绘制饼状图
    plt.legend()
    plt.savefig('../docs/pics/news_pie.jpg')  # 保存到本地文件夹，当前路径下
    doc.add_picture('../docs/pics/news_pie.jpg')  # 向报告中插入图片
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
    p.paragraph_format.first_line_indent = Inches(0.32)  # 首行缩进
    run = p.add_run('其中，积极倾向的舆论共' + str(pos_num) + '条，中性倾向的舆论共' + str(neu_num) + '条， 消极倾向的舆论共'
                    + str(neg_num) + '条。在这些舆论中，具备积极倾向且热度最高的舆论如下：')
    run.font.size = Pt(14)
    # 以表格形式筛选并显示热度最高积极舆论
    guba_pos_list = []
    for per in guba:
        if per[4] > 0:
            guba_pos_list.append(per)
        if len(guba_pos_list) > 4:
            break
    tb = doc.add_table(5, 5)
    tb.style = 'Medium Grid 3 Accent 1'
    tb.style.font.size = Pt(12)
    tb.cell(0, 1).width = Inches(6)
    heading_cells = tb.rows[0].cells
    heading_cells[0].text = '作者'
    heading_cells[1].text = '舆论内容'
    heading_cells[2].text = '发布时间'
    heading_cells[3].text = '浏览数'
    heading_cells[4].text = '评论数'
    for i in range(4):
        tb.cell(i + 1, 0).text = guba_pos_list[i][1]
        tb.cell(i + 1, 1).text = guba_pos_list[i][0]
        tb.cell(i + 1, 2).text = datetime.strftime(guba_pos_list[i][6], '%Y-%m-%d %H:%M:%S')
        tb.cell(i + 1, 3).text = str(guba_pos_list[i][2])
        tb.cell(i + 1, 4).text = str(guba_pos_list[i][3])
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
    p.paragraph_format.first_line_indent = Inches(0.32)  # 首行缩进
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run('在这些舆论中，具备消极倾向且热度最高的舆论如下：')
    run.font.size = Pt(14)
    # 以表格形式筛选并显示热度最高积极舆论
    guba_pos_list = []
    for per in guba:
        if per[4] < 0:
            guba_pos_list.append(per)
        if len(guba_pos_list) > 4:
            break
    tb = doc.add_table(5, 5)
    tb.style = 'Medium Grid 3 Accent 1'
    tb.style.font.size = Pt(12)
    tb.cell(0, 1).width = Inches(6)
    heading_cells = tb.rows[0].cells
    heading_cells[0].text = '作者'
    heading_cells[1].text = '舆论内容'
    heading_cells[2].text = '发布时间'
    heading_cells[3].text = '浏览数'
    heading_cells[4].text = '评论数'
    for i in range(4):
        tb.cell(i + 1, 0).text = guba_pos_list[i][1]
        tb.cell(i + 1, 1).text = guba_pos_list[i][0]
        tb.cell(i + 1, 2).text = datetime.strftime(guba_pos_list[i][6], '%Y-%m-%d %H:%M:%S')
        tb.cell(i + 1, 3).text = str(guba_pos_list[i][2])
        tb.cell(i + 1, 4).text = str(guba_pos_list[i][3])
    # 计算综合情感倾向分数
    c_s_sum = 0
    for per in guba:
        c_s = per[5]
        c_s_sum += c_s
    c_s_ave = round(c_s_sum / len(guba), 1)
    if c_s_ave > 0:
        c_att = '积极'
    elif c_s_ave < 0:
        c_att = '消极'
    else:
        c_att = '中性'
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
    p.paragraph_format.first_line_indent = Inches(0.32)  # 首行缩进
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run('将舆论数据增加热度权重后，得到的加权平均情感分数为' + str(c_s_ave) + '，整体呈现偏' + c_att + '倾向。')
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('三、上证指数相关舆论数据')
    run.font.color.rgb = RGBColor(60, 60, 60)
    run.font.size = Pt(20)
    run.font.bold = True
    # 获取并显示舆论内容
    status, guba = get_today_guba('sh000001')
    guba_num = len(guba)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
    p.paragraph_format.first_line_indent = Inches(0.32)  # 首行缩进
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(14)
    if status == 0:
        introduction = '今日共有上证指数相关舆情内容0篇，数据来源为东方财富股吧。'
    else:
        introduction = '今日共有上证指数相关舆情内容' + str(guba_num) + '篇, 数据来源为东方财富股吧，分析如下：'
    run = p.add_run(introduction)
    run.font.size = Pt(14)
    # 使用matplotlib绘制正负面新闻占比饼状图
    pos_num = 0
    neg_num = 0
    neu_num = 0
    for per in guba:
        score = per[4]
        if score > 0:
            pos_num += 1
        elif score < 0:
            neg_num += 1
        else:
            neu_num += 1
    plt.rcParams['font.family'] = ['Heiti TC']  # 设置正常显示中文
    plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号
    labels = ['积极', '中性', '消极']  # 定义饼的标签名称
    fraces = [pos_num, neu_num, neg_num]  # 显示百分比数据
    explode = [0.1, 0, 0]  # 饼图分离
    plt.axis('equal')  # 设置x,y的刻度一样，使其饼图为正圆
    plt.figure(figsize=(6, 2.5), dpi=150)
    cmap = plt.get_cmap("tab20c")  # 选择tab20c的色彩，
    color = cmap([6, 9, 1])  # 选择tab20c中的4种颜色
    plt.pie(x=fraces, labels=labels, autopct='%0.2f%%', colors=color, explode=explode, shadow=True)  # 绘制饼状图
    plt.legend()
    plt.savefig('../docs/pics/news_pie.jpg')  # 保存到本地文件夹，当前路径下
    doc.add_picture('../docs/pics/news_pie.jpg')  # 向报告中插入图片
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
    p.paragraph_format.first_line_indent = Inches(0.32)  # 首行缩进
    run = p.add_run('其中，积极倾向的舆论共' + str(pos_num) + '条，中性倾向的舆论共' + str(neu_num) + '条， 消极倾向的舆论共'
                    + str(neg_num) + '条。在这些舆论中，具备积极倾向且热度最高的舆论如下：')
    run.font.size = Pt(14)
    # 以表格形式筛选并显示热度最高积极舆论
    guba_pos_list = []
    for per in guba:
        if per[4] > 0:
            guba_pos_list.append(per)
        if len(guba_pos_list) > 4:
            break
    tb = doc.add_table(5, 5)
    tb.style = 'Medium Grid 3 Accent 1'
    tb.style.font.size = Pt(12)
    tb.cell(0, 1).width = Inches(6)
    heading_cells = tb.rows[0].cells
    heading_cells[0].text = '作者'
    heading_cells[1].text = '舆论内容'
    heading_cells[2].text = '发布时间'
    heading_cells[3].text = '浏览数'
    heading_cells[4].text = '评论数'
    for i in range(4):
        tb.cell(i + 1, 0).text = guba_pos_list[i][1]
        tb.cell(i + 1, 1).text = guba_pos_list[i][0]
        tb.cell(i + 1, 2).text = datetime.strftime(guba_pos_list[i][6], '%Y-%m-%d %H:%M:%S')
        tb.cell(i + 1, 3).text = str(guba_pos_list[i][2])
        tb.cell(i + 1, 4).text = str(guba_pos_list[i][3])
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
    p.paragraph_format.first_line_indent = Inches(0.32)  # 首行缩进
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run('在这些舆论中，具备消极倾向且热度最高的舆论如下：')
    run.font.size = Pt(14)
    # 以表格形式筛选并显示热度最高积极舆论
    guba_pos_list = []
    for per in guba:
        if per[4] < 0:
            guba_pos_list.append(per)
        if len(guba_pos_list) > 4:
            break
    tb = doc.add_table(5, 5)
    tb.style = 'Medium Grid 3 Accent 1'
    tb.style.font.size = Pt(12)
    tb.cell(0, 1).width = Inches(6)
    heading_cells = tb.rows[0].cells
    heading_cells[0].text = '作者'
    heading_cells[1].text = '舆论内容'
    heading_cells[2].text = '发布时间'
    heading_cells[3].text = '浏览数'
    heading_cells[4].text = '评论数'
    for i in range(4):
        tb.cell(i + 1, 0).text = guba_pos_list[i][1]
        tb.cell(i + 1, 1).text = guba_pos_list[i][0]
        tb.cell(i + 1, 2).text = datetime.strftime(guba_pos_list[i][6], '%Y-%m-%d %H:%M:%S')
        tb.cell(i + 1, 3).text = str(guba_pos_list[i][2])
        tb.cell(i + 1, 4).text = str(guba_pos_list[i][3])
    # 计算综合情感倾向分数
    c_s_sum = 0
    for per in guba:
        c_s = per[5]
        c_s_sum += c_s
    c_s_ave = round(c_s_sum / len(guba), 1)
    if c_s_ave > 0:
        c_att = '积极'
    elif c_s_ave < 0:
        c_att = '消极'
    else:
        c_att = '中性'
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
    p.paragraph_format.first_line_indent = Inches(0.32)  # 首行缩进
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run('将舆论数据增加热度权重后，得到的加权平均情感分数为' + str(c_s_ave) + '，整体呈现偏' + c_att + '倾向。')
    run.font.size = Pt(14)

    doc.save('../docs/舆情分析报告.docx')


if __name__ == '__main__':
    # print(get_today_guba('sh688180')[1])
    generate_report('sh688180')
    # print(code_to_name('sh688180'))
